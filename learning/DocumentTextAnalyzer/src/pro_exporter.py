"""
pro_exporter.py
PRO-Exportfunktionen: Text als DOCX oder PDF speichern.
"""

from datetime import datetime


def export_docx(text: str, save_path: str) -> tuple[bool, str]:
    """
    Speichert den übergebenen Text als DOCX-Datei.
    Gibt (erfolg, nachricht) zurück.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Seitenränder
        for section in doc.sections:
            section.top_margin    = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin   = Cm(3.0)
            section.right_margin  = Cm(2.5)

        # Kopfzeile
        hdr = doc.sections[0].header
        hdr_para = hdr.paragraphs[0]
        hdr_para.text = f"VT Document Text Converter  –  Exportiert am {datetime.now().strftime('%d.%m.%Y %H:%M')}"
        hdr_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hdr_para.runs[0].font.size = Pt(9)

        # Inhalt: Absätze splitten
        for line in text.split("\n"):
            para = doc.add_paragraph(line)
            if para.runs:
                para.runs[0].font.size = Pt(11)

        doc.save(save_path)
        return True, f"Als DOCX gespeichert: {save_path}"
    except ImportError:
        return False, "python-docx nicht installiert.\nBitte: pip install python-docx"
    except Exception as e:
        return False, f"DOCX-Export fehlgeschlagen: {e}"


def export_pdf(text: str, save_path: str) -> tuple[bool, str]:
    """
    Speichert den übergebenen Text als PDF-Datei.
    Gibt (erfolg, nachricht) zurück.
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        # Versuche DejaVu (breitere Unicode-Unterstützung) zu laden
        _register_fonts()

        doc = SimpleDocTemplate(
            save_path,
            pagesize=A4,
            leftMargin=3 * cm, rightMargin=2.5 * cm,
            topMargin=2.5 * cm, bottomMargin=2.5 * cm,
            title="VT Document Text Converter – Export",
            author="vt-solutions GmbH",
        )

        styles = getSampleStyleSheet()
        body_style = ParagraphStyle(
            "VTBody",
            parent=styles["Normal"],
            fontName=_get_font_name(),
            fontSize=11,
            leading=16,
            spaceAfter=4,
        )
        header_style = ParagraphStyle(
            "VTHeader",
            parent=styles["Normal"],
            fontName=_get_font_name(),
            fontSize=8,
            textColor=colors.grey,
            spaceAfter=6,
        )

        story = []

        # Kopfzeile
        ts = datetime.now().strftime("%d.%m.%Y %H:%M")
        story.append(Paragraph(
            f"VT Document Text Converter &nbsp;·&nbsp; Exportiert am {ts}",
            header_style,
        ))
        story.append(Spacer(1, 0.4 * cm))

        # Inhalt
        for line in text.split("\n"):
            safe = _escape(line) or "&nbsp;"
            story.append(Paragraph(safe, body_style))

        doc.build(story)
        return True, f"Als PDF gespeichert: {save_path}"

    except ImportError:
        return False, "reportlab nicht installiert.\nBitte: pip install reportlab"
    except Exception as e:
        return False, f"PDF-Export fehlgeschlagen: {e}"


# ------------------------------------------------------------------ #
#  Hilfsfunktionen                                                    #
# ------------------------------------------------------------------ #

_font_name = None


def _register_fonts() -> None:
    """Versucht DejaVuSans einzubinden (einmalig)."""
    global _font_name
    if _font_name is not None:
        return
    import os, sys
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    search_dirs = [
        os.path.join(os.environ.get("WINDIR", "C:\\Windows"), "Fonts"),
        "/usr/share/fonts/truetype/dejavu",
        "/usr/share/fonts",
    ]
    for d in search_dirs:
        candidate = os.path.join(d, "DejaVuSans.ttf")
        if os.path.isfile(candidate):
            try:
                pdfmetrics.registerFont(TTFont("DejaVuSans", candidate))
                _font_name = "DejaVuSans"
                return
            except Exception:
                pass
    _font_name = "Helvetica"


def _get_font_name() -> str:
    if _font_name is None:
        _register_fonts()
    return _font_name or "Helvetica"


def _escape(text: str) -> str:
    """Escapet XML-Sonderzeichen für ReportLab-Paragraphen."""
    return (
        text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
    )
