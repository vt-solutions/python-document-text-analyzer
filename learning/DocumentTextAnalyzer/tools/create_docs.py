"""
create_docs.py – Erzeugt die vollstaendige Lizenz-Dokumentation als Word-Dateien.
Ausgabe:
  docs/Technische_Dokumentation_Lizenz.docx
  docs/Benutzerdokumentation_Keygen.docx

Aufruf:
  python tools/create_docs.py
"""

import io
import sys
from pathlib import Path

# Projektroot
ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))
DOCS = ROOT / "docs"
DOCS.mkdir(exist_ok=True)

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, ImageDraw, ImageFont


# ═══════════════════════════════════════════════════════════════════════ #
#  Farben & Hilfsfunktionen                                              #
# ═══════════════════════════════════════════════════════════════════════ #

NAVY    = (10, 33, 58)
BLUE    = (30, 90, 160)
BLUE_L  = (220, 235, 255)
GRAY    = (90, 100, 115)
GRAY_L  = (240, 242, 245)
GREEN   = (22, 163, 74)
GREEN_L = (220, 250, 230)
RED     = (185, 28, 28)
RED_L   = (255, 220, 220)
ORANGE  = (180, 90, 20)
ORANGE_L = (255, 240, 210)
WHITE   = (255, 255, 255)
BLACK   = (20, 20, 30)

LOGO_PATH = ROOT / "assets" / "logo.png"


def _rgb(color):
    return RGBColor(*color)


def _set_cell_bg(cell, color_tuple):
    r, g, b = color_tuple
    hex_color = f"{r:02X}{g:02X}{b:02X}"
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _img_buf(img: Image.Image) -> io.BytesIO:
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf


def _try_font(size=14, bold=False):
    """Laed einen Windows-Systemfont oder faellt auf Standard zurueck."""
    paths = [
        "C:/Windows/Fonts/segoeui.ttf",
        "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibri.ttf",
    ]
    for p in paths:
        try:
            if bold:
                bp = p.replace(".ttf", "b.ttf").replace("segoeui", "segoeuib")
                try:
                    return ImageFont.truetype(bp, size)
                except Exception:
                    pass
            return ImageFont.truetype(p, size)
        except Exception:
            pass
    return ImageFont.load_default()


# ═══════════════════════════════════════════════════════════════════════ #
#  Diagramme (Pillow)                                                    #
# ═══════════════════════════════════════════════════════════════════════ #

def make_key_format_diagram() -> io.BytesIO:
    """Visuelles Schluesselformat-Diagramm."""
    W, H = 860, 240
    img = Image.new("RGB", (W, H), WHITE)
    d = ImageDraw.Draw(img)

    fn_big  = _try_font(26, bold=True)
    fn_med  = _try_font(18, bold=True)
    fn_sm   = _try_font(14)
    fn_tiny = _try_font(12)

    # Hintergrund-Banner
    d.rectangle([0, 0, W, 50], fill=NAVY)
    d.text((W // 2, 25), "Schluesselformat: VT-YYYY-EPPP-CCCC",
           fill=WHITE, font=fn_med, anchor="mm")

    # Segment-Boxen
    segs = [
        ("VT",   "Praefix",          "Hersteller-\nKennung",   BLUE,    BLUE_L,  60,  100),
        ("2026", "YYYY",             "Lizenzjahr\n(4 Ziffern)", GREEN,   GREEN_L, 180, 100),
        ("P001", "EPPP",             "E=Edition\nPPP=Seriennr.", ORANGE, ORANGE_L,340, 100),
        ("LD2G", "CCCC",             "HMAC-SHA256\nSignatur",   NAVY,    BLUE_L,  500, 100),
    ]

    for text, label, desc, fg, bg, x, y in segs:
        # Box
        d.rectangle([x, y, x + 120, y + 60], fill=bg, outline=fg, width=2)
        d.text((x + 60, y + 30), text, fill=fg, font=fn_big, anchor="mm")
        # Label unten
        d.text((x + 60, y + 75), label, fill=fg, font=fn_sm, anchor="mm")
        # Beschreibung
        for i, line in enumerate(desc.split("\n")):
            d.text((x + 60, y + 100 + i * 18), line, fill=GRAY, font=fn_tiny, anchor="mm")

    # Bindestrich-Trennungen
    for x in [120, 180, 300, 460]:
        d.text((x + 30, 130), "-", fill=GRAY, font=fn_big, anchor="mm")

    # Beispielzeile unten
    d.text((W // 2, 220), "Beispiel:  VT-2026-P001-LD2G  (PRO, Seriennummer 1)",
           fill=GRAY, font=fn_tiny, anchor="mm")

    return _img_buf(img)


def make_architecture_diagram() -> io.BytesIO:
    """Architektur-Uebersicht als Blockdiagramm."""
    W, H = 860, 420
    img = Image.new("RGB", (W, H), (248, 250, 252))
    d = ImageDraw.Draw(img)

    fn_bold = _try_font(16, bold=True)
    fn_med  = _try_font(14)
    fn_sm   = _try_font(12)

    # Titel
    d.rectangle([0, 0, W, 44], fill=NAVY)
    d.text((W // 2, 22), "Modul-Architektur: Lizenzsystem", fill=WHITE, font=fn_bold, anchor="mm")

    def box(x, y, w, h, title, lines, bg, fg):
        d.rectangle([x, y, x + w, y + h], fill=bg, outline=fg, width=2)
        d.text((x + w // 2, y + 18), title, fill=fg, font=fn_bold, anchor="mm")
        d.line([x + 8, y + 28, x + w - 8, y + 28], fill=fg, width=1)
        for i, ln in enumerate(lines):
            d.text((x + 10, y + 36 + i * 17), ln, fill=GRAY, font=fn_sm)

    # Schicht 1: App
    box(30,  60, 200, 90, "app.py", ["Hauptfenster", "UI-Logik", "Pro-Guard"], BLUE_L, BLUE)
    box(260, 60, 200, 90, "license_dialog.py", ["Aktivierungs-Dialog", "Trial-Start", "Status-Anzeige"], GREEN_L, GREEN)
    box(490, 60, 200, 90, "info_dialog.py", ["Info / Hilfe", "Versions-Info"], GRAY_L, GRAY)

    # Schicht 2: Manager
    box(145, 210, 260, 100, "license_manager.py", [
        "validate_key()", "activate()", "start_trial()",
        "has_feature()", "get_status_text()"
    ], ORANGE_L, ORANGE)

    # Schicht 3: Crypto
    box(60, 370, 220, 34, "license_crypto.py",
        [], (200, 220, 255), NAVY)
    d.text((60 + 110, 370 + 17),
           "HMAC-SHA256  |  generate_key()  |  get_machine_id()",
           fill=NAVY, font=fn_sm, anchor="mm")

    # Schicht 3: Storage
    box(330, 370, 200, 34, "license.json", [], GRAY_L, GRAY)
    d.text((330 + 100, 370 + 17),
           "%APPDATA%\\vt-solutions\\VTConverter",
           fill=GRAY, font=fn_sm, anchor="mm")

    # Schicht 3: Keygen
    box(570, 370, 220, 34, "tools/keygen.py", [], GREEN_L, GREEN)
    d.text((570 + 110, 370 + 17),
           "Internes Schluessel-Erzeugungstool",
           fill=GREEN, font=fn_sm, anchor="mm")

    # Pfeile (vereinfacht)
    def arrow(x1, y1, x2, y2):
        d.line([x1, y1, x2, y2], fill=GRAY, width=2)
        # Pfeilspitze
        dx, dy = x2 - x1, y2 - y1
        length = (dx**2 + dy**2) ** 0.5
        if length == 0:
            return
        ux, uy = dx / length, dy / length
        ax, ay = x2 - ux * 10 - uy * 5, y2 - uy * 10 + ux * 5
        bx, by = x2 - ux * 10 + uy * 5, y2 - uy * 10 - ux * 5
        d.polygon([(x2, y2), (ax, ay), (bx, by)], fill=GRAY)

    arrow(130, 150, 200, 210)
    arrow(360, 150, 280, 210)
    arrow(275, 310, 200, 370)
    arrow(275, 310, 380, 370)
    arrow(275, 310, 640, 370)

    return _img_buf(img)


def make_validation_flow() -> io.BytesIO:
    """Ablaufdiagramm der Schluessel-Validierung."""
    W, H = 480, 520
    img = Image.new("RGB", (W, H), WHITE)
    d = ImageDraw.Draw(img)

    fn_bold = _try_font(14, bold=True)
    fn_sm   = _try_font(12)

    d.rectangle([0, 0, W, 40], fill=NAVY)
    d.text((W // 2, 20), "Schluessel-Validierungsablauf", fill=WHITE, font=fn_bold, anchor="mm")

    def rbox(x, y, w, h, text, bg, fg):
        d.rectangle([x, y, x + w, y + h], fill=bg, outline=fg, width=2)
        for i, ln in enumerate(text.split("\n")):
            d.text((x + w // 2, y + h // 2 + (i - len(text.split('\n')) // 2) * 16),
                   ln, fill=fg, font=fn_sm, anchor="mm")

    def diamond(cx, cy, text):
        hw, hh = 120, 28
        d.polygon([
            (cx, cy - hh), (cx + hw, cy),
            (cx, cy + hh), (cx - hw, cy)
        ], fill=ORANGE_L, outline=ORANGE, width=2)
        d.text((cx, cy), text, fill=ORANGE, font=fn_sm, anchor="mm")

    def farr(x1, y1, x2, y2, label=""):
        d.line([x1, y1, x2, y2], fill=GRAY, width=2)
        d.polygon([(x2, y2), (x2 - 5, y2 - 8), (x2 + 5, y2 - 8)], fill=GRAY)
        if label:
            d.text(((x1 + x2) // 2 + 8, (y1 + y2) // 2), label, fill=GRAY, font=fn_sm)

    # Schritt 1
    rbox(140, 55, 200, 36, "Schluessel eingeben", BLUE_L, BLUE)
    farr(240, 91, 240, 120)

    # Schritt 2
    diamond(240, 150, "Format OK?")
    farr(240, 178, 240, 208)
    d.text((290, 155), "Nein", fill=RED, font=fn_sm)
    d.line([360, 150, 430, 150], fill=RED, width=2)
    rbox(370, 130, 100, 36, "Abgelehnt\n(Format)", RED_L, RED)

    # Schritt 3
    rbox(140, 208, 200, 36, "Editions-Code pruefen\n(P / S)", BLUE_L, BLUE)
    farr(240, 244, 240, 274)

    # Schritt 4
    rbox(140, 274, 200, 48, "HMAC-SHA256\nberechnen", ORANGE_L, ORANGE)
    farr(240, 322, 240, 352)

    # Schritt 5
    diamond(240, 380, "Signatur OK?")
    farr(240, 408, 240, 438)
    d.text((290, 382), "Nein", fill=RED, font=fn_sm)
    d.line([360, 380, 430, 380], fill=RED, width=2)
    rbox(370, 362, 100, 36, "Abgelehnt\n(Faelschung)", RED_L, RED)

    # Ergebnis
    rbox(130, 440, 220, 46, "Akzeptiert\nEdition: PRO / STANDARD", GREEN_L, GREEN)

    # Ja-Labels
    d.text((246, 130), "Ja", fill=GREEN, font=fn_sm)
    d.text((246, 362), "Ja", fill=GREEN, font=fn_sm)

    return _img_buf(img)


def make_keygen_flow() -> io.BytesIO:
    """Ablaufdiagramm der Schluessel-Erzeugung."""
    W, H = 480, 440
    img = Image.new("RGB", (W, H), WHITE)
    d = ImageDraw.Draw(img)

    fn_bold = _try_font(14, bold=True)
    fn_sm   = _try_font(12)

    d.rectangle([0, 0, W, 40], fill=GREEN)
    d.text((W // 2, 20), "Schluessel-Erzeugungsablauf (keygen.py)", fill=WHITE, font=fn_bold, anchor="mm")

    def rbox(x, y, w, h, text, bg, fg):
        d.rectangle([x, y, x + w, y + h], fill=bg, outline=fg, width=2)
        for i, ln in enumerate(text.split("\n")):
            offs = (i - len(text.split('\n')) // 2) * 16
            d.text((x + w // 2, y + h // 2 + offs), ln, fill=fg, font=fn_sm, anchor="mm")

    def farr(x1, y1, x2, y2):
        d.line([x1, y1, x2, y2], fill=GRAY, width=2)
        d.polygon([(x2, y2), (x2 - 5, y2 - 8), (x2 + 5, y2 - 8)], fill=GRAY)

    steps = [
        ("Edition + Jahr + Seriennr.\nwaehlen",              BLUE_L, BLUE),
        ("Editions-Code bestimmen\nP (PRO) | S (STANDARD)", BLUE_L, BLUE),
        ("Seriennr. in Base36\nkodieren  (0-9, A-Z)",       ORANGE_L, ORANGE),
        ("Payload zusammensetzen\nz. B. P001",               ORANGE_L, ORANGE),
        ("HMAC-SHA256 berechnen\nueber 'VT-2026-P001'",     (200, 220, 255), NAVY),
        ("Signatur auf 4 Zeichen\nkuerzen (mod 36^4)",       (200, 220, 255), NAVY),
        ("Schluessel ausgeben\nVT-2026-P001-LD2G",           GREEN_L, GREEN),
    ]

    y = 60
    for text, bg, fg in steps:
        rbox(90, y, 300, 44, text, bg, fg)
        y += 44
        if y < 60 + 44 * len(steps):
            farr(240, y, 240, y + 10)
        y += 10

    return _img_buf(img)


def make_feature_table_img() -> io.BytesIO:
    """Visuelle Feature-Tabelle."""
    cols = ["Feature", "Trial", "Standard", "PRO"]
    rows = [
        ("PDF-Extraktion",              True,  True,  True),
        ("Bild-OCR (JPG/PNG/TIFF)",     True,  True,  True),
        ("Office (DOCX/XLSX/PPTX)",     True,  True,  True),
        ("In Zwischenablage kopieren",  True,  True,  True),
        ("Als TXT speichern",           True,  True,  True),
        ("30-Tage-Testlauf",            True,  False, False),
        ("Ext. Formate TXT/ODT/RTF",    False, False, True),
        ("Export als DOCX",             False, False, True),
        ("Export als PDF",              False, False, True),
        ("Batch-Verarbeitung",          False, False, True),
        ("Textbearbeitung im Editor",   False, False, True),
    ]

    cw = [280, 90, 110, 90]
    rh = 32
    pad = 8
    W = sum(cw) + 2
    H = (len(rows) + 2) * rh + 2

    img = Image.new("RGB", (W, H), WHITE)
    d = ImageDraw.Draw(img)
    fn_hd = _try_font(14, bold=True)
    fn_sm = _try_font(13)

    # Header
    x = 0
    for i, (col, w) in enumerate(zip(cols, cw)):
        bg = NAVY if i == 0 else (GREEN if i == 3 else BLUE)
        d.rectangle([x, 0, x + w, rh], fill=bg)
        d.text((x + w // 2, rh // 2), col, fill=WHITE, font=fn_hd, anchor="mm")
        x += w

    # Datenzeilen
    for ri, (feat, t, s, p) in enumerate(rows):
        y = (ri + 1) * rh
        bg_row = WHITE if ri % 2 == 0 else GRAY_L
        d.rectangle([0, y, W, y + rh], fill=bg_row)
        # Feature-Name
        d.text((pad, y + rh // 2), feat, fill=BLACK, font=fn_sm, anchor="lm")
        # Haekchen
        x = cw[0]
        for val, w in zip([t, s, p], cw[1:]):
            mark = "  JA" if val else "  -"
            color = GREEN if val else RED
            d.text((x + w // 2, y + rh // 2), mark, fill=color, font=fn_sm, anchor="mm")
            x += w

    # Gitterlinien
    for ri in range(len(rows) + 2):
        y = ri * rh
        d.line([0, y, W, y], fill=(200, 205, 215), width=1)
    x = 0
    for w in cw:
        d.line([x, 0, x, H], fill=(200, 205, 215), width=1)
        x += w

    return _img_buf(img)


# ═══════════════════════════════════════════════════════════════════════ #
#  Word-Dokument-Helfer                                                  #
# ═══════════════════════════════════════════════════════════════════════ #

def _apply_base_styles(doc: Document):
    """Setzt Schrift + Seitenraender fuer das gesamte Dokument."""
    section = doc.sections[0]
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = "Segoe UI"
    style.font.size = Pt(10.5)
    style.font.color.rgb = _rgb(BLACK)


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = _rgb(NAVY)
    p.runs[0].font.size = Pt(18)
    p.runs[0].font.name = "Segoe UI"
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = _rgb(BLUE)
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.name = "Segoe UI"
    return p


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = _rgb(NAVY)
    p.runs[0].font.size = Pt(11.5)
    p.runs[0].font.name = "Segoe UI"
    return p


def body(doc, text, bold=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "Segoe UI"
    run.font.bold = bold
    if color:
        run.font.color.rgb = _rgb(color)
    return p


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9.5)
    run.font.color.rgb = _rgb((30, 60, 100))
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "E8EFF8")
    p._p.get_or_add_pPr().append(shd)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.6)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "Segoe UI"
    return p


def note_box(doc, text, color=BLUE_L, border=BLUE):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, color)
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Segoe UI"
    run.font.color.rgb = _rgb(NAVY)
    doc.add_paragraph()


def add_img(doc, buf: io.BytesIO, caption: str, width_cm=15):
    doc.add_picture(buf, width=Cm(width_cm))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = _rgb(GRAY)
    cap.runs[0].font.italic = True
    doc.add_paragraph()


def simple_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header-Zeile
    hdr = tbl.rows[0]
    for i, (cell, text) in enumerate(zip(hdr.cells, headers)):
        _set_cell_bg(cell, NAVY)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.bold = True
        run.font.color.rgb = _rgb(WHITE)
        run.font.size = Pt(10)
        run.font.name = "Segoe UI"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Datenzeilen
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        bg = WHITE if ri % 2 == 0 else GRAY_L
        for ci, (cell, text) in enumerate(zip(row.cells, row_data)):
            _set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            run.font.size = Pt(10)
            run.font.name = "Segoe UI"

    if col_widths:
        for row in tbl.rows:
            for cell, w in zip(row.cells, col_widths):
                cell.width = Cm(w)

    doc.add_paragraph()
    return tbl


def add_cover(doc, title, subtitle, version="1.0"):
    """Titelseite."""
    doc.add_paragraph()
    doc.add_paragraph()

    if LOGO_PATH.exists():
        doc.add_picture(str(LOGO_PATH), width=Cm(7))
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    run.font.name = "Segoe UI"
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = _rgb(NAVY)

    p2 = doc.add_paragraph()
    run2 = p2.add_run(subtitle)
    run2.font.name = "Segoe UI"
    run2.font.size = Pt(14)
    run2.font.color.rgb = _rgb(BLUE)

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    run3 = p3.add_run(
        f"VT Document Text Converter\nvt-solutions GmbH\n"
        f"Dokument-Version {version}  |  Stand: Juni 2026"
    )
    run3.font.name = "Segoe UI"
    run3.font.size = Pt(10)
    run3.font.color.rgb = _rgb(GRAY)

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════ #
#  TECHNISCHE DOKUMENTATION                                              #
# ═══════════════════════════════════════════════════════════════════════ #

def build_tech_doc():
    doc = Document()
    _apply_base_styles(doc)

    # ── Titelseite ─────────────────────────────────────────────────── #
    add_cover(doc,
              "Technische Dokumentation",
              "Lizenzmanagementsystem – VT Document Text Converter",
              version="1.0")

    # ── 1. Uebersicht ─────────────────────────────────────────────── #
    h1(doc, "1  Systemuebersicht")
    body(doc,
         "Das Lizenzmanagementsystem des VT Document Text Converters arbeitet vollstaendig "
         "offline und benoetigt keine Internetverbindung. Es basiert auf kryptografischer "
         "HMAC-SHA256-Schluessel-Validierung und unterscheidet drei Lizenz-Editionen.")

    doc.add_paragraph()
    add_img(doc, make_architecture_diagram(),
            "Abbildung 1 – Modul-Architektur des Lizenzsystems", width_cm=16)

    h2(doc, "1.1  Editionen")
    simple_table(doc,
        ["Edition", "Zweck", "Laufzeit", "Schluessel erforderlich"],
        [
            ["TRIAL",    "Testversion mit vollem Standard-Funktionsumfang", "30 Tage", "Nein"],
            ["STANDARD", "Produktiv-Lizenz mit Basisfunktionen",             "Dauerhaft", "Ja"],
            ["PRO",      "Alle Funktionen inkl. Export, Batch, Extraformate","Dauerhaft", "Ja"],
        ],
        col_widths=[3.5, 7.5, 3, 4]
    )

    # ── 2. Schluesselformat ────────────────────────────────────────── #
    h1(doc, "2  Schluesselformat")
    body(doc,
         "Jeder Lizenzschluessel folgt dem Format VT-YYYY-EPPP-CCCC. "
         "Die vier Segmente sind durch Bindestriche getrennt und enthalten "
         "ausschliesslich Grossbuchstaben und Ziffern (A-Z, 0-9).")

    add_img(doc, make_key_format_diagram(),
            "Abbildung 2 – Aufbau eines Lizenzschluessels", width_cm=16)

    simple_table(doc,
        ["Segment", "Zeichen", "Inhalt", "Beispiel"],
        [
            ["VT",   "2",  "Fester Herstellerprafix",          "VT"],
            ["YYYY", "4",  "Lizenzjahr (4-stellig)",            "2026"],
            ["EPPP", "4",  "E = Editions-Code (P/S), PPP = Seriennr. (Base36)", "P001 = PRO Nr.1"],
            ["CCCC", "4",  "HMAC-SHA256-Signatur (Base36, mod 36^4)", "LD2G"],
        ],
        col_widths=[2, 2.5, 9, 4.5]
    )

    h2(doc, "2.1  Editions-Codes")
    simple_table(doc,
        ["Code", "Edition", "Beschreibung"],
        [
            ["P", "PRO",      "Schaltet alle PRO-Features frei"],
            ["S", "STANDARD", "Basisfunktionen ohne PRO-Extras"],
        ],
        col_widths=[2, 4, 12]
    )

    h2(doc, "2.2  Seriennummer (PPP)")
    body(doc,
         "Die drei Zeichen nach dem Editions-Code (PPP) bilden die Seriennummer "
         "in Base36-Kodierung (Ziffern 0-9, Buchstaben A-Z). "
         "Damit sind 36^3 = 46.656 Seriennummern pro Edition und Jahr moeglich.")
    simple_table(doc,
        ["Seriennr.", "Base36-Wert", "Beispiel-Schluessel"],
        [
            ["1",    "001", "VT-2026-P001-LD2G"],
            ["10",   "00A", "VT-2026-P00A-XXXX"],
            ["100",  "02S", "VT-2026-S02S-9H2R"],
            ["1295", "ZZZ", "VT-2026-SZZZ-XXXX"],
        ],
        col_widths=[4, 4, 10]
    )

    # ── 3. Kryptografische Validierung ────────────────────────────── #
    h1(doc, "3  Kryptografische Schluessel-Validierung")
    body(doc,
         "Die Signatur (letztes Segment CCCC) wird mit HMAC-SHA256 unter Verwendung eines "
         "im Quellcode eingebetteten, XOR-maskierten Geheimschluessels berechnet. "
         "Ohne Kenntnis dieses Schluessels kann keine gueltige Signatur erzeugt werden.")

    add_img(doc, make_validation_flow(),
            "Abbildung 3 – Ablauf der Schluessel-Validierung", width_cm=10)

    h2(doc, "3.1  Algorithmus")
    body(doc, "Schritt-fuer-Schritt-Berechnung der Signatur:")
    code_block(doc,
        "payload  = 'P001'                           # Editions-Code + Seriennummer\n"
        "year     = '2026'\n"
        "message  = f'VT-{year}-{payload}'           # = 'VT-2026-P001'\n"
        "digest   = HMAC-SHA256(SECRET_KEY, message) # 32-Byte-Hash\n"
        "value    = int.from_bytes(digest[:4], 'big') % 36^4  # 0 – 1.679.615\n"
        "signature = base36_encode(value, width=4)   # z. B. 'LD2G'"
    )

    h2(doc, "3.2  Signatur-Staerke")
    simple_table(doc,
        ["Parameter", "Wert"],
        [
            ["Hash-Algorithmus",          "HMAC-SHA256"],
            ["Signaturlaenge",            "4 Zeichen (Base36)"],
            ["Moegliche Signaturen",      "36^4 = 1.679.616"],
            ["Faelschungswahrscheinlichkeit", "~ 1 : 1.679.616 pro zufaelligem Versuch"],
            ["Geheimschluessel-Laenge",   "26 Bytes (XOR-maskiert gespeichert)"],
            ["Online-Verbindung noetig",  "Nein – vollstaendig offline"],
        ],
        col_widths=[8, 10]
    )

    h2(doc, "3.3  Geheimschluessel (Obfuskierung)")
    body(doc,
         "Der Geheimschluessel ist in license_crypto.py als XOR-maskiertes Byte-Array "
         "gespeichert. Der Maskierungs-Schluessel (0x4F) ist ebenfalls im Code hinterlegt; "
         "der effektive Schutz ergibt sich aus der Kombination beider Werte, die fuer einen "
         "Angreifer ohne Quellcode-Zugang nicht trivial zu rekonstruieren sind.")
    note_box(doc,
        "WICHTIG: Den Geheimschluessel niemals aendern, nachdem Lizenzen ausgegeben wurden. "
        "Alle bestehenden Schluessel wuerden sonst unguelig. Fuer neue Produktversionen "
        "koennen separate Schluessel-Generationen mit neuem Jahrescode (YYYY) eingefuehrt werden.",
        color=ORANGE_L, border=ORANGE)

    # ── 4. Feature-Management ─────────────────────────────────────── #
    h1(doc, "4  Feature-Management")
    body(doc,
         "Alle Programmfunktionen werden ueber das zentrale Feature-Gate-System "
         "in license_manager.py kontrolliert. Die Funktion has_feature(feature) "
         "gibt True zurueck, wenn die aktive Edition das angeforderte Feature enthaelt.")

    add_img(doc, make_feature_table_img(),
            "Abbildung 4 – Feature-Uebersicht nach Edition", width_cm=16)

    h2(doc, "4.1  Feature-Bezeichner")
    simple_table(doc,
        ["Bezeichner", "Beschreibung", "Edition"],
        [
            ["basic_extraction", "PDF-, OCR- und Office-Extraktion",       "TRIAL, STD, PRO"],
            ["copy",             "In Zwischenablage kopieren",              "TRIAL, STD, PRO"],
            ["save_txt",         "Als TXT-Datei speichern",                 "TRIAL, STD, PRO"],
            ["extra_formats",    ".txt, .odt, .rtf als Eingabe-Formate",   "PRO"],
            ["export_docx",      "Text als DOCX exportieren",               "PRO"],
            ["export_pdf",       "Text als PDF exportieren",                "PRO"],
            ["batch",            "Batch-Verarbeitung mehrerer Dateien",     "PRO"],
            ["edit_text",        "Direkte Textbearbeitung im Editor",       "PRO"],
        ],
        col_widths=[4, 8, 4.5]
    )

    h2(doc, "4.2  Neues Feature hinzufuegen")
    body(doc, "Um ein neues PRO-Feature hinzuzufuegen, sind genau zwei Schritte noetig:")
    bullet(doc, "In _FEATURES['PRO'] in license_manager.py den Bezeichner eintragen.")
    bullet(doc, "Im App-Code: if lm.has_feature('neues_feature'): ... aufrufen.")
    code_block(doc,
        "# license_manager.py – _FEATURES erweitern:\n"
        "_FEATURES = {\n"
        "    ...\n"
        "    'PRO': {\n"
        "        ...,\n"
        "        'mein_neues_feature',   # <-- hier eintragen\n"
        "    },\n"
        "}"
    )

    # ── 5. Hardware-Bindung ───────────────────────────────────────── #
    h1(doc, "5  Hardware-Bindung (Machine Binding)")
    body(doc,
         "Die Hardware-Bindung verhindert, dass ein Lizenzschluessel auf mehreren "
         "Computern gleichzeitig genutzt wird. Sie ist standardmaessig deaktiviert "
         "(MACHINE_BINDING_ENABLED = False in license_manager.py).")

    h2(doc, "5.1  Aktivierung")
    code_block(doc,
        "# license_manager.py, Zeile 38:\n"
        "MACHINE_BINDING_ENABLED = True   # auf True setzen, um Bindung zu aktivieren"
    )
    note_box(doc,
        "Hinweis: Bei aktivierter Hardware-Bindung wird beim ersten Aktivieren die Geraete-ID "
        "gespeichert. Bei Hardware-Wechsel (z. B. neues Mainboard) muss der Nutzer "
        "erneut aktivieren und ggf. den Support kontaktieren.",
        color=ORANGE_L, border=ORANGE)

    h2(doc, "5.2  Geraete-ID-Berechnung")
    body(doc, "Die Geraete-ID setzt sich zusammen aus:")
    bullet(doc, "Motherboard-Seriennummer (Windows WMIC – stabil auch nach OS-Neuinstallation)")
    bullet(doc, "Hostname als Fallback, falls Mainboard-Seriennummer nicht lesbar ist")
    code_block(doc,
        "# get_machine_id() in license_crypto.py gibt z. B. zurueck:\n"
        "'A3F82C1D9E4B7F06'  # 16-stelliger SHA-256-Hash"
    )

    # ── 6. Lizenzdatei ───────────────────────────────────────────── #
    h1(doc, "6  Lizenzdatei (license.json)")
    body(doc,
         "Die Lizenzinformationen werden lokal unter "
         "%APPDATA%\\vt-solutions\\VTConverter\\license.json gespeichert. "
         "Nachfolgend das JSON-Schema fuer eine aktivierte PRO-Lizenz:")
    code_block(doc,
        '{\n'
        '  "customer":     "Max Mustermann",\n'
        '  "license_key":  "VT-2026-P001-LD2G",\n'
        '  "activated":    true,\n'
        '  "edition":      "PRO",\n'
        '  "activated_at": "2026-06-14T10:30:00.000000",\n'
        '  "machine_id":   "A3F82C1D9E4B7F06"  // nur wenn MACHINE_BINDING_ENABLED=True\n'
        '}'
    )
    body(doc, "Trial-Version (Beispiel):")
    code_block(doc,
        '{\n'
        '  "customer":      "Testbenutzer",\n'
        '  "license_key":   "TRIAL",\n'
        '  "activated":     false,\n'
        '  "edition":       "TRIAL",\n'
        '  "trial_started": "2026-06-14T10:00:00.000000",\n'
        '  "trial_ends":    "2026-07-14T10:00:00.000000"\n'
        '}'
    )

    # ── 7. Modul-Referenz ────────────────────────────────────────── #
    h1(doc, "7  Modul-API-Referenz")

    h2(doc, "7.1  license_crypto.py")
    simple_table(doc,
        ["Funktion", "Parameter", "Rueckgabe", "Beschreibung"],
        [
            ["validate_key(key)", "str", "(bool, str)", "Kryptografische Pruefung"],
            ["generate_key(ed, yr, sn)", "str, int, int", "str", "Schluessel erzeugen"],
            ["get_machine_id()", "–", "str", "16-Hex-Geraete-ID"],
        ],
        col_widths=[5, 4, 3.5, 5.5]
    )

    h2(doc, "7.2  license_manager.py")
    simple_table(doc,
        ["Funktion", "Rueckgabe", "Beschreibung"],
        [
            ["validate_key(key)",    "(bool, str)", "Format + Krypto-Pruefung"],
            ["activate(name, key)",  "(bool, str)", "Lizenz aktivieren + speichern"],
            ["start_trial()",        "dict",         "Trial-Lizenz starten"],
            ["deactivate()",         "None",         "Lizenzdatei loeschen"],
            ["has_feature(feat)",    "bool",         "Feature-Gate-Abfrage"],
            ["is_activated()",       "bool",         "Prueft aktive Lizenz + ggf. Maschine"],
            ["is_trial_expired()",   "bool",         "Trial-Ablauf pruefen"],
            ["get_edition()",        "str",          "Aktuelle Edition"],
            ["get_status_text()",    "str",          "Anzeige-Text fuer Footer"],
            ["check_machine_binding()", "bool",      "Geraete-ID-Vergleich"],
        ],
        col_widths=[5.5, 3.5, 9]
    )

    # ── 8. Sicherheitshinweise ────────────────────────────────────── #
    h1(doc, "8  Sicherheitshinweise")
    note_box(doc,
        "Das System ist auf Offline-Nutzung ausgelegt. Der Schutz basiert auf der "
        "Geheimhaltung des HMAC-Schluessels. Bei Verdacht auf Kompromittierung sollte "
        "der Schluessel geaendert und alle Lizenzen neu ausgestellt werden.",
        color=ORANGE_L, border=ORANGE)

    bullet(doc, "Geheimschluessel niemals in Git-Repositories oder Logs veroeffentlichen.")
    bullet(doc, "Die EXE-Datei mit PyInstaller --onefile oder vergleichbar ausliefern (kein reiner Quellcode).")
    bullet(doc, "Keygen-Tool (tools/keygen.py) ausschliesslich intern verwenden und nicht ausliefern.")
    bullet(doc, "Seriennummern intern verwalten, um Doppelvergabe zu vermeiden (siehe Benutzerdokumentation).")
    bullet(doc, "Fuer hoehere Sicherheit: MACHINE_BINDING_ENABLED = True setzen.")

    path = DOCS / "Technische_Dokumentation_Lizenz.docx"
    doc.save(str(path))
    print(f"[OK] Erstellt: {path}")
    return path


# ═══════════════════════════════════════════════════════════════════════ #
#  BENUTZERDOKUMENTATION                                                 #
# ═══════════════════════════════════════════════════════════════════════ #

def build_user_doc():
    doc = Document()
    _apply_base_styles(doc)

    add_cover(doc,
              "Benutzerdokumentation",
              "Lizenzverwaltung & Key-Generator – Interne Anleitung",
              version="1.0")

    # Inhaltsverzeichnis (statisch)
    h1(doc, "Inhaltsverzeichnis")
    toc_entries = [
        ("1", "Einfuehrung & Uebersicht"),
        ("2", "Voraussetzungen"),
        ("3", "Key-Generator: Schritt-fuer-Schritt"),
        ("4", "Schluessel verwalten (Seriennummern-Liste)"),
        ("5", "Schluessel pruefen"),
        ("6", "EXE-Datei mit Schluessel ausliefern"),
        ("7", "Endbenutzer: Lizenz aktivieren"),
        ("8", "Endbenutzer: Trial-Version starten"),
        ("9", "Haeufige Fragen & Fehler"),
    ]
    for nr, titel in toc_entries:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.3)
        run = p.add_run(f"  {nr}   {titel}")
        run.font.name = "Segoe UI"
        run.font.size = Pt(11)
    doc.add_page_break()

    # ── 1. Einfuehrung ────────────────────────────────────────────── #
    h1(doc, "1  Einfuehrung & Uebersicht")
    body(doc,
         "Dieses Dokument beschreibt den internen Prozess zur Verwaltung, "
         "Erzeugung und Auslieferung von Lizenzschluesseln fuer den "
         "VT Document Text Converter. Die Anleitung richtet sich an Mitarbeiter "
         "von vt-solutions GmbH, die Kunden beliefern.")

    add_img(doc, make_feature_table_img(),
            "Abbildung 1 – Funktionsumfang nach Edition", width_cm=16)

    note_box(doc,
        "Trial = 30 Tage, identisch mit Standard-Funktionsumfang.\n"
        "Standard = Dauerlizenz, Basisfunktionen.\n"
        "PRO = Dauerlizenz, alle Funktionen freigeschaltet.",
        color=BLUE_L, border=BLUE)

    # ── 2. Voraussetzungen ────────────────────────────────────────── #
    h1(doc, "2  Voraussetzungen")
    body(doc, "Folgendes muss auf dem Rechner des Mitarbeiters installiert sein:")
    simple_table(doc,
        ["Komponente", "Version", "Hinweis"],
        [
            ["Python",           "3.11 oder hoeher", "Im Projektordner als .venv verfuegbar"],
            ["python-docx",      "1.2.0",             "Bereits im .venv enthalten"],
            ["VT Projektordner", "develop-Branch",    "C:\\...\\DocumentTextAnalyzer"],
        ],
        col_widths=[4.5, 3.5, 10]
    )

    body(doc, "Der Projektordner-Pfad:", bold=True)
    code_block(doc, "C:\\Users\\<Name>\\PythonProjects\\learning\\DocumentTextAnalyzer")

    # ── 3. Key-Generator ──────────────────────────────────────────── #
    h1(doc, "3  Key-Generator: Schritt-fuer-Schritt")

    h2(doc, "3.1  Wie wird der Keygen aufgerufen?")
    body(doc,
         "Der Key-Generator wird als Kommandozeilen-Tool ausgefuehrt. "
         "Oeffnen Sie ein Terminal (PowerShell oder Eingabeaufforderung) "
         "im Projektordner und aktivieren Sie die virtuelle Umgebung:")

    code_block(doc,
        "# Schritt 1: In den Projektordner wechseln\n"
        "cd C:\\Users\\<Name>\\PythonProjects\\learning\\DocumentTextAnalyzer\n\n"
        "# Schritt 2: Virtuelle Umgebung aktivieren (PowerShell)\n"
        ".venv\\Scripts\\Activate.ps1\n\n"
        "# Alternativ: CMD\n"
        ".venv\\Scripts\\activate.bat"
    )

    h2(doc, "3.2  PRO-Schluessel erzeugen")
    body(doc, "Beispiel: 10 PRO-Schluessel fuer das Jahr 2026 ab Seriennummer 1:")
    code_block(doc,
        "python tools\\keygen.py --edition PRO --year 2026 --serial 1 --count 10"
    )
    body(doc, "Ausgabe:")
    code_block(doc,
        "+-------------------------------------------------------+\n"
        "|   VT Document Text Converter  -  Key Generator        |\n"
        "|   vt-solutions GmbH  |  NUR INTERNER GEBRAUCH         |\n"
        "+-------------------------------------------------------+\n"
        "  Edition : PRO\n"
        "  Jahr    : 2026\n"
        "  Anzahl  : 10  (Start-Seriennummer: 1)\n\n"
        "  Schluessel\n"
        "  ------------------------------------------------\n"
        "  VT-2026-P001-LD2G   (#1)\n"
        "  VT-2026-P002-7OJH   (#2)\n"
        "  VT-2026-P003-P1LV   (#3)\n"
        "  VT-2026-P004-0QOX   (#4)\n"
        "  VT-2026-P005-L9JE   (#5)\n"
        "  ...                 ..."
    )

    add_img(doc, make_keygen_flow(),
            "Abbildung 2 – Interner Ablauf der Schluessel-Erzeugung", width_cm=10)

    h2(doc, "3.3  STANDARD-Schluessel erzeugen")
    code_block(doc,
        "python tools\\keygen.py --edition STANDARD --year 2026 --serial 100 --count 5"
    )
    body(doc, "Ausgabe-Beispiel:")
    code_block(doc,
        "  VT-2026-S02S-9H2R   (#100)\n"
        "  VT-2026-S02T-P455   (#101)\n"
        "  VT-2026-S02U-GLQL   (#102)\n"
        "  VT-2026-S02V-JURC   (#103)\n"
        "  VT-2026-S02W-UKLI   (#104)"
    )

    h2(doc, "3.4  Alle Parameter im Ueberblick")
    simple_table(doc,
        ["Parameter", "Pflicht", "Standard", "Beschreibung", "Beispiel"],
        [
            ["--edition", "Ja",  "–",    "Lizenz-Edition",              "PRO oder STANDARD"],
            ["--year",    "Nein","2026",  "Lizenzjahr (4-stellig)",      "--year 2027"],
            ["--serial",  "Nein","1",     "Erste Seriennummer",          "--serial 50"],
            ["--count",   "Nein","1",     "Anzahl zu erzeugender Schl.", "--count 20"],
            ["--verify",  "Nein","–",     "Einzelnen Schluessel pruefen","--verify VT-2026-P001-LD2G"],
        ],
        col_widths=[2.8, 2, 2.5, 6.5, 4.2]
    )

    note_box(doc,
        "Beachten Sie: Das Jahr im Schluessel (YYYY) dient nur als Teil der Signatur-Berechnung "
        "und stellt kein automatisches Ablaufdatum dar. Ein VT-2026-...-Schluessel laeuft "
        "NICHT am 31.12.2026 ab. Er ist dauerhaft gueltig.",
        color=BLUE_L, border=BLUE)

    # ── 4. Seriennummern verwalten ────────────────────────────────── #
    h1(doc, "4  Schluessel verwalten (Seriennummern-Liste)")
    body(doc,
         "Um Doppelvergaben zu vermeiden, muss intern eine Liste gepflegt werden, "
         "welche Seriennummern bereits vergeben wurden. Die folgende Tabelle zeigt "
         "ein empfohlenes Verwaltungsschema (z. B. als Excel-Datei):")

    simple_table(doc,
        ["Seriennr.", "Schluessel", "Edition", "Kunde", "Datum", "Status"],
        [
            ["1",   "VT-2026-P001-LD2G", "PRO",      "Mustermann GmbH",  "14.06.2026", "Aktiv"],
            ["2",   "VT-2026-P002-7OJH", "PRO",      "Test AG",          "14.06.2026", "Aktiv"],
            ["100", "VT-2026-S02S-9H2R", "STANDARD", "Klein & Partner",  "14.06.2026", "Aktiv"],
            ["101", "VT-2026-S02T-P455", "STANDARD", "–",                "–",          "Reserviert"],
        ],
        col_widths=[2.5, 5, 3, 4.5, 3, 2]
    )

    body(doc, "Wichtige Regeln bei der Seriennummern-Verwaltung:", bold=True)
    bullet(doc, "PRO und STANDARD haben getrennte Seriennummer-Reihen (P001, S001).")
    bullet(doc, "Pro Bestellung wird eine neue, noch nicht vergebene Seriennummer verwendet.")
    bullet(doc, "Stornierte Lizenzen koennen NICHT wiederverwendet werden.")
    bullet(doc, "Seriennummern-Liste regelmaessig sichern (Backup).")
    bullet(doc, "Fuer jedes neue Jahr (--year 2027) kann der Zaehler neu bei 1 beginnen.")

    note_box(doc,
        "Empfehlung: PRO-Schluessel ab #1 vergeben, STANDARD-Schluessel ab #1000 vergeben. "
        "Dadurch ist aus dem Schluessel sofort ersichtlich, welche Edition vorliegt.",
        color=GREEN_L, border=GREEN)

    # ── 5. Schluessel pruefen ─────────────────────────────────────── #
    h1(doc, "5  Schluessel pruefen")
    body(doc,
         "Mit dem Parameter --verify kann jederzeit geprueft werden, "
         "ob ein Schluessel gueltig ist und welcher Edition er angehoert:")

    code_block(doc,
        "# Gueltiger PRO-Schluessel\n"
        "python tools\\keygen.py --verify VT-2026-P001-LD2G\n\n"
        "  Schluessel : VT-2026-P001-LD2G\n"
        "  Status     : GUELTIG\n"
        "  Edition    : PRO\n\n"
        "# Ungueltig (falsche Signatur)\n"
        "python tools\\keygen.py --verify VT-2026-P001-XXXX\n\n"
        "  Schluessel : VT-2026-P001-XXXX\n"
        "  Status     : UNGUELTIG"
    )

    add_img(doc, make_validation_flow(),
            "Abbildung 3 – Ablauf der Schluessel-Validierung", width_cm=10)

    # ── 6. EXE-Datei mit Schluessel ausliefern ────────────────────── #
    h1(doc, "6  EXE-Datei mit Schluessel ausliefern")

    h2(doc, "6.1  Was wird ausgeliefert?")
    body(doc,
         "Der Endbenutzer erhaelt zwei Bestandteile, die separat uebermittelt werden. "
         "Niemals den Schluessel fest in die EXE einbetten:")

    simple_table(doc,
        ["Bestandteil", "Format", "Wie uebermitteln"],
        [
            ["Setup-Datei / EXE", ".exe oder .zip",
             "Download-Link, USB, E-Mail-Anhang"],
            ["Lizenzschluessel",  "Textzeile (VT-2026-...)",
             "Separate E-Mail oder Lieferschein"],
        ],
        col_widths=[5, 3, 10]
    )

    h2(doc, "6.2  Schritt-fuer-Schritt: Auslieferungsprozess")

    steps_delivery = [
        ("Schritt 1", "Seriennummer auswaehlen",
         "Naechste freie Seriennummer aus der Verwaltungsliste entnehmen."),
        ("Schritt 2", "Schluessel erzeugen",
         "python tools\\keygen.py --edition PRO --year 2026 --serial 42 --count 1"),
        ("Schritt 3", "Schluessel pruefen",
         "python tools\\keygen.py --verify VT-2026-P0142-XXXX"),
        ("Schritt 4", "Seriennummer als 'Vergeben' markieren",
         "In der Verwaltungsliste eintragen: Kunde, Datum, Schluessel."),
        ("Schritt 5", "EXE-Installer an Kunden senden",
         "Via E-Mail-Anhang oder Downloadlink (kein Schluessel in der EXE)."),
        ("Schritt 6", "Schluessel separat mitteilen",
         "Schluessel in einer eigenen E-Mail oder auf dem Lieferschein."),
        ("Schritt 7", "Kunde aktiviert selbst",
         "Kunde startet EXE, traegt Namen + Schluessel ein, klickt 'Aktivieren'."),
    ]

    for nr, titel, beschr in steps_delivery:
        tbl = doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        nr_cell = tbl.cell(0, 0)
        _set_cell_bg(nr_cell, NAVY)
        nr_cell.width = Cm(3)
        nr_run = nr_cell.paragraphs[0].add_run(nr)
        nr_run.font.bold = True
        nr_run.font.color.rgb = _rgb(WHITE)
        nr_run.font.size = Pt(10)
        nr_run.font.name = "Segoe UI"

        txt_cell = tbl.cell(0, 1)
        _set_cell_bg(txt_cell, GRAY_L)
        p = txt_cell.paragraphs[0]
        r1 = p.add_run(titel + "\n")
        r1.font.bold = True
        r1.font.size = Pt(10)
        r1.font.name = "Segoe UI"
        r1.font.color.rgb = _rgb(NAVY)
        r2 = p.add_run(beschr)
        r2.font.size = Pt(10)
        r2.font.name = "Segoe UI"
        doc.add_paragraph()

    h2(doc, "6.3  E-Mail-Vorlage fuer Schluessel-Uebermittlung")
    note_box(doc,
        "Betreff: Ihr VT Document Text Converter Lizenzschluessel\n\n"
        "Sehr geehrte/r [Kundenname],\n\n"
        "vielen Dank fuer Ihren Kauf des VT Document Text Converters [Edition].\n\n"
        "Ihr persoenlicher Lizenzschluessel lautet:\n\n"
        "    VT-2026-P001-LD2G\n\n"
        "So aktivieren Sie Ihre Lizenz:\n"
        "1. Starten Sie die Anwendung (VTConverter.exe).\n"
        "2. Geben Sie Ihren Namen oder Firmennamen ein.\n"
        "3. Tragen Sie den obigen Schluessel ein.\n"
        "4. Klicken Sie auf 'Lizenz aktivieren'.\n\n"
        "Bei Fragen: support@vt-solutions.de\n\n"
        "Mit freundlichen Gruessen\nvt-solutions GmbH",
        color=BLUE_L, border=BLUE)

    h2(doc, "6.4  Was muss NICHT ausgeliefert werden?")
    body(doc, "Folgende Dateien sind intern und duerfen den Kunden NICHT erreichen:", bold=True)
    simple_table(doc,
        ["Datei / Ordner", "Begruendung"],
        [
            ["tools/keygen.py",             "Enthaelt die Schluessel-Erzeugungslogik"],
            ["src/licensing/license_crypto.py", "Enthaelt den (obfuskierten) Geheimschluessel"],
            ["src/licensing/license_manager.py","Kernlogik des Lizenzsystems"],
            [".venv/",                       "Interne Python-Umgebung"],
            ["Quellcode generell (.py)",     "Schutzmassnahme: nur EXE ausliefern"],
        ],
        col_widths=[7, 11]
    )
    note_box(doc,
        "Die EXE-Datei wird mit PyInstaller erstellt und enthaelt bereits den kompilierten "
        "Python-Code inklusive aller Abhaengigkeiten. Kunden erhalten nur die EXE, "
        "niemals den Quellcode.",
        color=ORANGE_L, border=ORANGE)

    # ── 7. Endbenutzer: Aktivierung ───────────────────────────────── #
    h1(doc, "7  Endbenutzer: Lizenz aktivieren")
    body(doc,
         "Diese Anleitung kann direkt an den Endkunden weitergegeben werden.")

    h2(doc, "7.1  Erststart der Anwendung")
    body(doc,
         "Beim ersten Start des VT Document Text Converters erscheint automatisch "
         "das Lizenzaktivierungs-Fenster. Es erscheint auch nach Ablauf der Testversion.")

    body(doc, "Das Fenster besteht aus zwei Bereichen:", bold=True)
    bullet(doc, "Links: Lizenz aktivieren (fuer Kunden mit Schluessel)")
    bullet(doc, "Rechts: 30-Tage-Testversion starten (kein Schluessel noetig)")

    h2(doc, "7.2  Lizenz aktivieren")
    body(doc, "Vorgehen Schritt fuer Schritt:")

    steps_act = [
        ("1", "Benutzernamen eingeben",
         "Tragen Sie Ihren vollstaendigen Namen oder den Firmennamen ein.\n"
         "Beispiel: Max Mustermann oder Mustermann GmbH"),
        ("2", "Lizenzschluessel eingeben",
         "Tragen Sie den Schluessel genauso ein, wie er in der E-Mail steht.\n"
         "Format: VT-2026-XXXX-XXXX  (mit Bindestrichen)\n"
         "Gross-/Kleinschreibung spielt keine Rolle."),
        ("3", "Auf 'Lizenz aktivieren' klicken",
         "Das Programm prueft den Schluessel sofort offline.\n"
         "Bei Erfolg erscheint 'Lizenz erfolgreich aktiviert!' und die\n"
         "Anwendung startet automatisch."),
    ]

    for nr, titel, text in steps_act:
        tbl = doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        nc = tbl.cell(0, 0)
        _set_cell_bg(nc, BLUE)
        nc.width = Cm(1.5)
        nr_run = nc.paragraphs[0].add_run(nr)
        nr_run.font.bold = True
        nr_run.font.color.rgb = _rgb(WHITE)
        nr_run.font.size = Pt(16)
        nr_run.font.name = "Segoe UI"
        nc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        tc = tbl.cell(0, 1)
        _set_cell_bg(tc, BLUE_L)
        p = tc.paragraphs[0]
        r1 = p.add_run(titel + "\n")
        r1.font.bold = True
        r1.font.size = Pt(10.5)
        r1.font.name = "Segoe UI"
        r1.font.color.rgb = _rgb(NAVY)
        r2 = p.add_run(text)
        r2.font.size = Pt(10)
        r2.font.name = "Segoe UI"
        doc.add_paragraph()

    h2(doc, "7.3  Fehlermeldungen bei der Aktivierung")
    simple_table(doc,
        ["Fehlermeldung", "Ursache", "Loesung"],
        [
            ["Bitte Benutzernamen eingeben",
             "Feld leer gelassen",
             "Namen eintragen"],
            ["Ungültiger Lizenzschluessel",
             "Schluessel falsch eingegeben oder manipuliert",
             "Schluessel nochmals aus der E-Mail kopieren"],
            ["Erwartet: VT-YYYY-XXXX-XXXX",
             "Format-Fehler (z. B. fehlender Bindestrich)",
             "Format genau beachten: 4 Gruppen mit Bindestrich"],
            ["Lizenz – Anderer PC",
             "Hardware-Bindung aktiv, anderer Computer",
             "Support kontaktieren: support@vt-solutions.de"],
        ],
        col_widths=[5, 5.5, 7.5]
    )

    # ── 8. Trial ──────────────────────────────────────────────────── #
    h1(doc, "8  Endbenutzer: Trial-Version starten")
    body(doc,
         "Falls noch kein Lizenzschluessel vorhanden ist, kann die "
         "30-Tage-Testversion gestartet werden.")

    bullet(doc, "Im Aktivierungsfenster rechts auf 'Testversion starten' klicken.")
    bullet(doc, "Kein Name, kein Schluessel erforderlich.")
    bullet(doc, "Die Testversion laeuft 30 Tage ab dem ersten Start.")
    bullet(doc, "Nach Ablauf erscheint das Aktivierungsfenster erneut.")
    bullet(doc, "Die Testversion entspricht funktional der Standard-Edition.")

    note_box(doc,
        "Nach Ablauf der 30 Tage koennen keine Dateien mehr verarbeitet werden, "
        "bis eine gueltige Lizenz aktiviert wird. Der bisher extrahierte Text geht "
        "nicht verloren.",
        color=ORANGE_L, border=ORANGE)

    # ── 9. FAQ ────────────────────────────────────────────────────── #
    h1(doc, "9  Haeufige Fragen & Fehler")

    faqs = [
        ("Kann ich einen Schluessel auf zwei PCs nutzen?",
         "Nein, wenn Hardware-Bindung aktiv ist. Ja, wenn nicht. Standardmaessig ist die "
         "Bindung deaktiviert. Klaeren Sie das mit Ihrer Lizenzrichtlinie."),
        ("Was passiert, wenn ich Windows neu installiere?",
         "Die Lizenzdatei geht verloren. Der Kunde muss den Schluessel erneut eintragen. "
         "Der Schluessel selbst bleibt gueltig."),
        ("Wo wird die Lizenz gespeichert?",
         "%APPDATA%\\vt-solutions\\VTConverter\\license.json\n"
         "(Typisch: C:\\Users\\<Name>\\AppData\\Roaming\\vt-solutions\\VTConverter\\license.json)"),
        ("Ich habe den Schluessel verloren – was tun?",
         "Die Verwaltungsliste bei vt-solutions GmbH auflegen und den Schluessel erneut ausgeben."),
        ("Kann ein Schluessel des Jahres 2026 auch 2027 genutzt werden?",
         "Ja. Das Jahr im Schluessel (YYYY) ist Teil der Signatur, kein Ablaufdatum."),
        ("Wie erkenne ich, ob ein Schluessel PRO oder Standard ist?",
         "Am fuenften Zeichen (Editions-Code): P = PRO, S = STANDARD.\n"
         "Beispiel: VT-2026-P001-... = PRO, VT-2026-S001-... = STANDARD."),
        ("Kann ich den Keygen auch auf einem anderen Rechner ausfuehren?",
         "Ja, solange Python + die Abhaengigkeiten installiert sind und "
         "der Geheimschluessel identisch ist. Empfehlung: Nur auf einem dedizierten "
         "internen Rechner ausfuehren."),
        ("Was wenn ein Kunde behauptet, sein Schluessel sei ungueltig?",
         "Schluessel mit --verify pruefen. Wenn ungueltig: neuen erzeugen. "
         "Wenn gueltig: Kunde bittet den Schluessel nochmals genau einzutippen (Tippfehler?)."),
    ]

    for frage, antwort in faqs:
        h3(doc, f"F: {frage}")
        body(doc, f"A: {antwort}")
        doc.add_paragraph()

    path = DOCS / "Benutzerdokumentation_Keygen.docx"
    doc.save(str(path))
    print(f"[OK] Erstellt: {path}")
    return path


# ═══════════════════════════════════════════════════════════════════════ #
#  MAIN                                                                  #
# ═══════════════════════════════════════════════════════════════════════ #

if __name__ == "__main__":
    print("\nErzeuge Dokumentation ...\n")
    p1 = build_tech_doc()
    p2 = build_user_doc()
    print(f"\nFertig. Beide Dateien liegen in:\n  {DOCS}\n")
